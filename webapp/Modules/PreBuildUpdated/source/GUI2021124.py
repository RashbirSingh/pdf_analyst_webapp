import os, time, uuid, pprint, json, copy, xlsxwriter, fitz, math, datetime, ocrmypdf
import PIL.Image, PIL.ImageTk
from mergedeep import merge, Strategy
from copy import deepcopy

from tkinter import *
from tkinter import messagebox, ttk, filedialog, N, S, E, W
from tkcolorpicker import askcolor

## used by Word export
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

## used by Zoom function
from tkinter import ALL, EventType

from HiLightModular20200602 import *
from HiColorsDict3 import *
from arrangeAliases6b import arrangeAliases

##arrangeAliases6 flattens the tree, causes other issues
## arrangeAliases7 tries to solve for honorifics..

#Test these imports - are they used?
import pandas as pd
import  PySimpleGUI as sg
from itertools import groupby
from ast import literal_eval
from operator import itemgetter
from operator import itemgetter
from tkDragDrop import TreeSetUp

#not used at present
#import pyperclip,
#import numpy as np


globaldebug = False
debug=False or globaldebug
debug2=False or globaldebug
debug3=False or globaldebug
debug4 = False or globaldebug
debug5 = False or globaldebug
debug6 = False or globaldebug
debug7 = False or globaldebug

#Preferences
#TODO preferences screen

exportHyperlinks = True
DEFAULT_CULTURE = Culture.English 
LOCALES = ['en-AU'] #affects parsing of Dates
##LOCALES = ['en'] #if US dates format wanted 
pathsep = "/" #OS dependent
if debug: print("pathsep: ",pathsep)
outfldr = "" #optional output subfolder
ModesDict = {
        "Highlight": "Highlight",
        "Underline": "Underline",
        "Rectangle": "Rect",
        "Squiggly underline": "Squiggly"
    }

result = {}
results = {}
docs = {}
d = {}
DocDict = {}
DocDictList = {}
##textPagesDict = {}
textSentencesDict = {}
edited_result = {}
file = ""
file_list = []
file_sel_list = []
## pageinFocus is the same as the page number, which is offset by +1 from the zero-based index of textPagesDict
pageinFocus= 1
fileinFocus = ""
root=None
textWidget = None

defaultColorDict = copy.deepcopy(InvColorDictLabelstoColors)
checkVar = {}
radioVar = {}


class CustomizeHL(object):
## Preferably, this would be a separat module    

    root = None

    global InvColorDictLabelstoColors, pprint, ModesDict
    
    def __init__(self, default):

        self.top = Toplevel(CustomizeHL.root)

        frm = Frame(self.top, borderwidth=4, relief='ridge')

        self.createWidgets(False)

        frm.title("pdfAnalyst - Customize highlights")

    def createWidgets(self, default):

        global InvColorDictLabelstoColors, ModesDict

        buttonNames = list(d.keys())
        buttonNames.append("DEFAULT")
        buttonDic = {}
        checkbuttonDic = {}
        scaleDic ={}
        radioDic = {}
        buttonColors = {}
        parent = self.top

        def on_OK():
            for i,key in enumerate(buttonNames[:-1]):
                a = checkVar[key]
                b = scaleDic[key]
                c = radioVar[key]
                InvColorDictLabelstoColors[key][3] = a.get()
                InvColorDictLabelstoColors[key][2] = b.get()
                InvColorDictLabelstoColors[key][1] = str(c.get())
            self.top.destroy()

        def on_save():
            for i,key in enumerate(buttonNames[:-1]):
                a = checkVar[key]
                b = scaleDic[key]
                c = radioVar[key]
                InvColorDictLabelstoColors[key][3] = a.get()
                InvColorDictLabelstoColors[key][2] = b.get()
                InvColorDictLabelstoColors[key][1] = str(c.get())
           
            with open('data.json', 'w') as fp:
                json.dump(InvColorDictLabelstoColors, fp)

        def on_load():
            global InvColorDictLabelstoColors          
            with open('data.json', 'r') as f:
                InvColorDictLabelstoColors = json.load(f)
##            pprint(InvColorDictLabelstoColors)
            self.createWidgets(False)
        
        def hex_to_rgb(value):
            value = value.lstrip('#')
            lv = len(value)
            return tuple(int(value[i:i + lv // 3], 16) for i in range(0, lv, lv // 3))

        def rgb255_to_rgb1(tpl):
            return (round(((tpl[0]+1)/256)-(1/256), 2),\
                    round(((tpl[1]+1)/256)-(1/256), 2), round(((tpl[2]+1)/256)-(1/256),2))
        
        def change_color(key):

            newcolor = askcolor((127, 127, 127), root)[1]
            buttonColors[key].set(newcolor)
            InvColorDictLabelstoColors[key][0] = rgb255_to_rgb1(hex_to_rgb(buttonColors[key].get()))
            buttonDic[key].config(background=buttonColors[key].get())
            self.top.tkraise()
##            self.top.tkraise(aboveThis=root)

        if default:
            InvColorDictLabelstoColors = copy.deepcopy(defaultColorDict)
             
        for key in d:

            rgb = (int(InvColorDictLabelstoColors[key][0][0]*255),
                   int(InvColorDictLabelstoColors[key][0][1]*255),
                   int(InvColorDictLabelstoColors[key][0][2]*255))
            keycolor = f'#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}'
            buttonColors[key] = StringVar()
            buttonColors[key].set(keycolor)
            buttonDic[key] = Button(
                    parent,
                    text = key,
                    width = 20,
                    background = buttonColors[key].get(),
                    command= lambda passName=key: change_color(passName)
                    )

            checkVar[key] = IntVar()
            checkbuttonDic[key] = Checkbutton(parent,text = "", variable = checkVar[key], onvalue=1,  offvalue=0)
            if InvColorDictLabelstoColors[key][3] == 1:
                checkVar[key].set(1)
            if checkVar[key] == 1: checkbuttonDic[key].select()

            scaleDic[key] = Scale(parent, from_=0, to=1, resolution=0.01, orient=HORIZONTAL)
            scaleDic[key].set(InvColorDictLabelstoColors[key][2])

            radioVar[key] = StringVar()
            radioVar[key].set(InvColorDictLabelstoColors[key][1]) # initialize
            radioDic[key] = {}
            for item in ModesDict:
                radioDic[key][item] = Radiobutton(parent, text=item, variable=radioVar[key], value=ModesDict[item])
           
        buttonDic["DEFAULT"] = Button(
                parent,
                text = "RESTORE DEFAULTS",
                width = 20,
                background = "white",
                command = lambda passName="DEFAULT": self.createWidgets(True)
                )
        
        parent.grid()
        First = False
        
        info_label0_0 = Label(parent, fg='red')
        info_label0_0.grid(row=0, column=0, pady=0, sticky=N+E+W)
        info_label0_0.config(text="Choose colors\nfor categories")
        info_label0_1 = Label(parent, fg='red')
        info_label0_1.grid(row=0, column=1, pady=0, sticky=N+E+W)
        info_label0_1.config(text="Checked categories\nwill be highlighted")
        info_label0_2 = Label(parent, fg='red')
        info_label0_2.grid(row=0, column=2, pady=0, sticky=S)
        info_label0_2.config(text="Adjust opacity")
        info_label0_3 = Label(parent, fg='red')
        info_label0_3.grid(row=0, column=3, columnspan = 4, pady=0, sticky = S)
        info_label0_3.config(text="Choose highlighting style")

        i = 0
        for i,key in enumerate(buttonNames[:-1]):
            buttonDic[key].grid(row=i+1,column=0)
            
        for i,key in enumerate(buttonNames[:-1]):
            checkbuttonDic[key].grid(row=i+1,column=1, sticky = W+E)
            scaleDic[key].grid(row=i+1,column=2, sticky = W+E)
            counter = 0
            for item in ModesDict:
                radioDic[key][item].grid(row=i+1, column = 3 + counter)
                counter += 1
##            InvColorDictLabelstoColors[key][3] = checkVar[key]

        buttonDic["DEFAULT"].grid(row=i+2,column=0, columnspan = 7, sticky = W+E)    

        save_button = Button(parent, text="Save Settings")
        save_button.grid(row=i+3,column = 0, columnspan = 7, sticky = W+E)
        save_button['command'] = on_save

        load_button = Button(parent, text="Load Settings")
        load_button.grid(row=i+4,column = 0, columnspan = 7, sticky = W+E)
        load_button['command'] = on_load

        quit_button = Button(parent, text="OK")
        quit_button.grid(row=i+5,column = 0, columnspan = 7, sticky = W+E)
        quit_button['command'] = on_OK


def contain(item, listbox):
    iscontain = item in listbox.get(0, "end")
    return iscontain

def import_files(eff, listbox, label):
    global file_list, debug
    ls = list(filedialog.askopenfilenames(title="Select PDF(s) to analyse", filetypes=[('PDF Files', '.pdf')]))
    for item in ls:
        if item not in file_list:
            file_list.append(item)

    if debug: print(type(file_list))
    for file in file_list:
        if not contain(file, listbox):
            listbox.insert(END, file)
    file_num = len(file_list)
    if  file_num > 0:
        info = "%d files in sandbox" % (file_num)
    else:
        info = "No %s files in sandbox" % ("")
    label.config(text=info)
    if debug: print(type(file_list))
    return file_list

def sort_list(eff, listbox):
    file_list.sort()
    listbox.delete(0, END)
    for file in file_list:
        listbox.insert(END, file)

def remove_files(eff, listbox, label):
    global file_list, debug3
    file_list2 = []
    if debug3: print(file_list)
    ls = listbox.curselection()
    if debug3: print(ls)
    if ls == (): messagebox.showerror("Error", "Choose a file to remove from sandbox")
    if debug3: print(len(file_list), ls)
    offset = 0
    for index in ls:
        del file_list[(index-offset):(index-offset+1)]
        offset += 1
    listbox.delete(0, END)
    for file in file_list:
        listbox.insert(END, file)
    file_num = len(file_list)
    if  file_num > 0:
        info = "%d files in sandbox" % (file_num)
    else:
        info = "No %s files in sandbox" % ("")
    label.config(text=info)
    return file_list

def remove_all_files(eff, listbox, label):
    global file_list, debug
    file_list = []
    listbox.delete(0, END)
    return file_list
    label.config(text="0 files in sandbox")

def get_list(eff=None, listbox = object, label = object):

    global debug
    if debug: print("get_list")
    try:
        file_list = selpath(listbox)
        return file_list
    except:
        label.config(text="Please select one or more files on the list1")
        return []

def selpath(listbox = object):
    global file_list
    sel = listbox.curselection()
    return [file_list[int(x)] for x in sel]

def analyse_file(eff = None, listbox=object, listbox2 = object, label=object, label2 = object, label3 = object, label4 = object, tree=object, root=object, scalePage = object, button = object, button2 = object):
    global file, docs, results, debug, file_list, file_sel_list, d, DocDict, DocDictList,textSentencesDict, InvColorDictLabelstoColors
    lst=[]
    if debug: print("analyse_file", listbox, label)

##
    
    try:
        sel = listbox.curselection()
        if debug: print(listbox.curselection())
        
        for i in sel:
            lst.append(file_list[i])
            
        if debug: print(lst)
        label3.config(text = "Analysing "+str(lst))
        if not tuple(lst) in results:
            results[tuple(lst)] = Highlight_Analyse(lst, InvColorDictLabelstoColors, False, False, False, label3, False)
        d = results[tuple(lst)][0]
##        print(d)
##        if debug4: pprint(d)
        DocDict = results[tuple(lst)][1]
##        print(DocDict)
##        if debug4: pprint(DocDict)
        DocDictList = results[tuple(lst)][2]
##        print(DocDictList)
##        textPagesDict = results[tuple(lst)][3]
        if debug4: pprint(DocDictList)
        textSentencesDict = results[tuple(lst)][3]
        print(textSentencesDict)
## TO DO deal with LineNumbers as [4] in results        
##        print(root, listbox2, label2, label4, textSentencesDict, scalePage)
        tree = CreateTree(root, listbox2, label2, label4, textSentencesDict, scalePage, button, button2)
        PopulateTree(tree, '', d, label3)
        if debug5: print(d)
        
    except NameError:
        label.config(text="Please select a file in the Sandbox.")

##    unconditional version of above can be useful for debugging
##    sel = listbox.curselection()
##    if debug: print(listbox.curselection())
##    
##    for i in sel:
##        lst.append(file_list[i])
##        
##    if debug: print(lst)
##
##    if not tuple(lst) in results:
##        results[tuple(lst)] = Highlight_Analyse(lst, InvColorDictLabelstoColors, False, False, False, label3, False)
##    d = results[tuple(lst)][0]
##    if debug4: pprint(d)
##    DocDict = results[tuple(lst)][1]
##    if debug4: pprint(DocDict)
##    DocDictList = results[tuple(lst)][2]
####        textPagesDict = results[tuple(lst)][3]
##    if debug4: pprint(DocDictList)
##    textSentencesDict = results[tuple(lst)][3]
#### TO DO deal with LineNumbers as [4] in results        
####        print(root, listbox2, label2, label4, textSentencesDict, scalePage)
##    tree = CreateTree(root, listbox2, label2, label4, textSentencesDict, scalePage, button, button2)
##    PopulateTree(tree, '', d, label3)
##    if debug5: print(d)

    file_sel_list = lst
    if debug: print("file_sel_lst: ",file_sel_list)

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def set_col_widths(table, dic):
    widths = []
    idx = 0
    for key in dic:
        widths.append(dic[key])
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Cm(width)

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def Export2Word(eff = None, label=object):
    global file, time, debug7, results, file_sel_list, exportHyperlinks
    timestr = time.strftime("%Y%m%d-%H%M%S")
    d = results[tuple(file_sel_list)][0]
    pth = file_sel_list[0]

    saveDocDict2Word(d, pth, timestr, file_sel_list, exportHyperlinks, debug7)
    label.config(text='Saved as Chronology_'+timestr+'.docx')

def rearrangeDate(string, sep):
    yr = string[0:4]
    mo = string[5:7]
    da = string[8:10]
    if LOCALES == ['en']:
        return (mo+sep+da+sep+yr)
    else:
        return (da+sep+mo+sep+yr)

def saveDocDict2Word(records:dict, input_pdf:str, timestr:str, file_sel_list:list, exportHyperlinks: bool, debug: bool):

    #Exports the DATE records as a formatted Chronology in Word. The columns include parsed date, unparsed (raw) date, source, page, sentence, context.

    document = Document()
    
    sections = document.sections
    
    margin = 2.54

    sep = "/"

##    format_ = '%Y-%m-%d'
##    format_ = '%Y-%m-%d'

    for section in sections:
        # change orientation to landscape
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

    document.add_heading('Combined chronology', 0)

    document.add_heading('List of files analysed', level=1)
    
    for file in file_sel_list:
        p = document.add_paragraph(
        file, style='List Bullet'
        )
##        add_hyperlink(p, text=str(file), url='http://www.demo.com', color=None, underline=True)

    labelDict = {"Date":2.5, "Text": 2.5, "Source":5, "Page": 0.7, "Sent.": 0.7, "Context": 10}
##    labelDict = {"Date":2.5, "EndDate":2.5, "Text": 2.5, "Source":2.5, "Page": 0.7, "Sent.": 0.7, "Context": 10}
    ##    labelDict = {"Term": 25, "File":25, "Page": 7, "Sent.": 7, "Context": 80, "Hyperlink": 50}

    document.add_heading('Chronology', level=1)
    
    table = document.add_table(rows=1, cols=len(labelDict))
##    print(table.autofit)
##    table.allow_autofit
##    print(table.autofit)
    
    hdr_cells = table.rows[0].cells
    make_rows_bold(table.rows[0])

    n = 0
    
    for key in labelDict:
        hdr_cells[n].text = key
        n += 1
        
    set_repeat_table_header(table.rows[0])

    for key in records:
        if key == "DATE":
            for key2 in records["DATE"]:
                
                for tupl in records["DATE"][key2]:
                    print("tupl: ",tupl)
                    row_cells = table.add_row().cells   

                    fiddled = False
                    if key2[:1] == "'":
                        key2 = guessCentury(key2[1:])
                        fiddled = True
                    d = MS_Recognize(key2)
                    dp = str(dateparser.parse(key2, locales=LOCALES))[0:10]

                    if fiddled == True:
                        key2 = "'"+key2[2:]
                    if not d == None:
                        try:
                            d = eval(d)
                            debug = False
                            if debug:
                                print("d: ",d)
                                print("dp: ",dp)
##                            print('d["values"][0]["type"]:',d["values"][0]["type"])

                            if str(d["values"][0]["type"]) == "date":
##                                print("True")
                                date1 = d["values"][0]["value"]
                                date1_time = rearrangeDate(date1, sep)
##                                date1_time = datetime.datetime.strptime(dp, format_)
##                                print("date1_time: ", date1_time)
##                                row_cells[0].text = "Demo"
                                row_cells[0].text = str(date1_time)[:10]

                            elif d["values"][0]["type"] == "daterange":
                                date1 = d["values"][0]["start"]
                                date1_time = rearrangeDate(date1, sep)
##                                datetime.datetime.strptime(dp, format_)
##                                date1_time = datetime.datetime.strptime(date1, '%Y-%m-%d')
                                row_cells[0].text = str(date1_time)[:10]
##                                row_cells[0].text = "Demo"
##                                row_cells[0].text = date1_time
                                
                         
                        except:
                            0

#version that does not include an End-date" field

                        row_cells[1].text = key2
                        n = 2
                        for item in tupl[:-1]:
                            row_cells[n].text = str(item)
                            n += 1

#version that includes an End-date" field

##                                date2 = d["values"][0]["end"]
##                                date2_time = rearrangeDate(date2, sep)
####                                date2_time = datetime.datetime.strptime(date2, format_)
####                                date2_time = datetime.datetime.strptime(date2, '%Y-%m-%d')
##                                row_cells[1].text = str(date2_time)[:10]
####                                row_cells[1].text = "Demo"
####                                row_cells[1].text = date2_time                            
##   
                            
##                        row_cells[2].text = key2
##                        n = 3
##                        for item in tupl[:-1]:
##                            row_cells[n].text = str(item)
##                            n += 1

##                        row_cells[n].text = str(tupl[-1])
##                        hyperlink = add_hyperlink(row_cells[n].text, str(tupl[-1]), str(tupl[-1]), None, True)
####                        hyperlink = p.add_hyperlink(document, text='Google', url='http://google.com')
                            
    set_col_widths(table, labelDict)
    filename = 'Chronology'+timestr+'.docx'
    dry = os.path.split(input_pdf)[0]
    newpath = os.path.join(dry,filename)
    document.save(newpath)
    

def ExportDicttoExcelUVO(eff = None, label=object):
    global file, time, debug, results, file_sel_list, exportHyperlinks
    timestr = time.strftime("%Y%m%d-%H%M%S")
    d = results[tuple(file_sel_list)][0]
    pth = file_sel_list[0]
    dry = os.path.split(pth)[0]
##    try:
    saveDict2ExcelUniqueValsOnly(d, dry+pathsep+outfldr+pathsep+"Analysis_", timestr, file_sel_list, exportHyperlinks, debug7)
    label.config(text="Saved as "+"Analysis_"+timestr+'UVO.xlsx')
##    except:
##        label.config(text="Export failed, please try other file(s)")

def ExportDetailstoExcel(eff = None, label=object):
    
    global file, time, debug7, results, file_sel_list, exportHyperlinks
    timestr = time.strftime("%Y%m%d-%H%M%S")
    d = results[tuple(file_sel_list)][0]
    pth = file_sel_list[0]
    dry = os.path.split(pth)[0]
##    try:
    if True:
        saveDocDict2Excel(d, dry+pathsep+outfldr+pathsep+"Analysis_", timestr, file_sel_list, exportHyperlinks, debug7)
        label.config(text="Saved as "+outfldr+os.path.sep+'Analysis_'+timestr+'DATA.xlsx')
##    except:
##        label.config(text="Export failed, please try other file(s)")

def ExportDetailstoWord(eff = None, label=object):
    
    global file, time, debug7, results, file_sel_list, exportHyperlinks
    timestr = time.strftime("%Y%m%d-%H%M%S")

    d = results[tuple(file_sel_list)][0]
    pth = file_sel_list[0]
    dry = os.path.split(pth)[0]
##    try:
    if True:
        saveDocDict2Word(d, dry+pathsep+outfldr+pathsep+"Analysis_", timestr, file_sel_list, exportHyperlinks, debug7)
        label.config(text="Saved as "+outfldr+os.path.sep+'Analysis_'+timestr+'DATA.docx')
##    except:
##        label.config(text="Export failed, please try other file(s)")

def ExporttoPDF(eff = None, label=object):
    global time, debug6, results, file_sel_list, outfldr
    if debug: print("in ExporttoPDF")
    timestr = time.strftime("%Y%m%d-%H%M%S")
    pathnew, p0, p1 = "", "", ""
    try:
        for file in file_sel_list:
    ######             print ("254")
            d2 = results[tuple(file_sel_list)][2]
    ##            pprint(d2)
            p0 = os.path.split(file)[0]
            p1 = os.path.split(file)[1]
    ##            doc = markup(p1, d2[file], InvColorDictLabelstoColors, debug7)
            doc = markup(file, d2[file], InvColorDictLabelstoColors, debug6)
    ##            saveDocToPDF(doc, p0, p1, outfldr, timestr, debug)
            pathnew = p0+pathsep+outfldr+pathsep+p1[:-4]+" "+timestr+" HI.pdf"
            doc.save(pathnew)
            doc.close()
        label.config(text="Saved with filename(s) ending in '"+timestr+"_HI.pdf'")
    except:
        label.config(text="Export failed, please try other file(s)"+str(file_sel_list))

def ExtractImages(eff = None, listbox= object, label = object):
## Saves images embedded in PDF in new pix dictionary
## images below50 * 50 are not saved
    global time, file_sel_list, debug
    n = 0
    file_sel_list = get_list(None, listbox, label)
    for j in file_sel_list:
        pth = j
        dry = os.path.split(pth)[0]
        fn = os.path.split(pth)[1]
        
        if not os.path.exists(dry+"/pix"):
            os.makedirs(dry+"/pix")

        doc = fitz.open(j)
        for i in range(len(doc)):
            for img in doc.getPageImageList(i):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                print (pix.n)
    ##            if pix.n < 5:
    ##                pix.writePNG("p%s-%s.png" % (i, xref))
    ##            else:
                if 1 == 1:
                    try:
                        pix1 = fitz.Pixmap(fitz.csRGB,pix)
                        print ("width: ", pix1.width," height: ",pix1.height)
                        if pix1.width > 50 and pix1.height > 50:
                            pix1.writePNG(dry+"/pix/"+fn[:-4]+"p."+str(i)+" #"+str(xref)+".png")
                            n += 1
                            print(fn[:-4]+"p."+str(i)+" #"+str(xref)+".png      "+ str(pix1.size))

                        else:
                            0
    ##                        pix1.writeJPG(dry+"/"+fn[:-4]+"p."+str(i)+" #"+str(xref)+".jpg")
                    except:
                        0
    ##                    print("p."+str(i)+" #"+str(xref))
    ##                pix1.writePNG(dry+"/"+"p%s-%s.png" % (i, xref))
                    pix1 = None
                pix = None
    if n == 0:
        label.config(text="No pictures saved.")
    else:
        label.config(text=str(n)+" picture(s) saved in subdirectory/ies named 'pix'")

def PDF2Images(eff = None, listbox= object, label = object):
## Saves each page of PDF as a separate image in new pix directory/ies
   
    global time, file_sel_list, debug
    f = 0
    
    file_sel_list = get_list(None, listbox, label)
    for j in file_sel_list:
        pth = j
        dry = os.path.split(pth)[0]
        fn = os.path.split(pth)[1]
        
        if not os.path.exists(dry+"/pix"):
            os.makedirs(dry+"/pix")

        doc = fitz.open(j)
        for page in doc:
            pix = page.getPixmap()
            pix.writeImage(dry+"/pix/"+fn[:-4]+" p."+str(page.number)+".png")
            pix = None

        f += 1
    if f == 0:
        label.config("No file saved.")
    else:
        label.config(text=str(f)+" file(s) saved as separate .png files in subdirectory/ies named 'pix'")

def DeleteHighlights(eff = None, listbox= object, label = object):

    global time, file_sel_list, debug7
    HLsFound = False
    failureMsg = msg = msg2 = ""
    if debug7: print("in DeleteHighlights")
    file_sel_list = get_list(None, listbox, label)

    timestr = time.strftime("%Y%m%d-%H%M%S")

    for file in file_sel_list:
        try:
            p0 = os.path.split(file)[0]
            p1 = os.path.split(file)[1]
            res = deleteHLs(file, debug)
            doc = res[0]
            HLsFound = HLsFound or res[1]
            if res[1]:
                pathnew = p0+pathsep+outfldr+pathsep+p1[:-4]+" "+"CLN.pdf"
                doc.save(pathnew)
                doc.close()
            else:
                failureMsg = res[2] + ", " + failureMsg 
        except:
            0
    if HLsFound:
        msg = "Cleaned file(s) saved with filename(s) ending in '_CLN.pdf'."
    if failureMsg != "":
        if failureMsg.endswith(", "): failureMsg = failureMsg[:-2] 
        msg2 = "No annotations found in these files: "+failureMsg+"."
    label.config(text=msg+msg2)
        
def deleteHLs(input_pdf, debug):
    HLsFound = False
    failureMsg = ""
    pdf = fitz.open(input_pdf)
    l, order = [], 0
    context = ""
    contentDict = {}
    for ixpage, page in enumerate(pdf):
        if debug: print (page)
        annot = page.firstAnnot

        while annot:

            HLsFound = True
            annot = page.deleteAnnot(annot)

    if HLsFound == False: failureMsg = os.path.split(input_pdf)[1] 
    return pdf, HLsFound, failureMsg
       
def close_children(parent):
    for child in tree.get_children(parent):
        close_children(child)

        
def CreateTree(root, textWidget, label, label2, textSentencesDict, scalePage, button, button2):

    def handleOpenEvent(event, parent, string):
        open_children(parent, string)

    def open_children(parent, string):
        tree.item(parent, open=True)
        for child in tree.get_children(string):
            open_children(parent, child)
    
    def handleCloseEvent2(event, tree, string):
        for child in tree.get_children(string):
            tree.item(child, open = False)

    def handleOpenEvent2(event, tree, string):
        for child in tree.get_children(string):
            tree.item(child, open = True)

##    opened = False
    tree = ttk.Treeview(root, columns = ('#0', '#1', '#2', '#3', '#4', '#5'))
    s = ttk.Style()
    s.configure('Treeview', rowheight=30)
##    , selectmode="extended"
    minwidth = tree.column('#0', option='minwidth')
    tree.column('#0', width=20)
    tree.heading('#0', text = 'Named entities', anchor = W)
    tree.column('#1', minwidth=0,width=10)
    tree.heading('#1', text='File', anchor=W)
##    tree.column('#2', minwidth=0,width=10)
##    tree.heading('#2', text='File', anchor=W)

    tree.column('#2', minwidth=0,width=1)
    tree.heading('#2', text='Page', anchor=W)
    tree.column('#3', minwidth=0,width=1)
    tree.heading('#3', text='Sentence', anchor=W)
    tree.column('#4', minwidth=0,width=150)
    tree.heading('#4', text='Context', anchor=W)
    tree.column('#5', minwidth=0,width=150)
    tree.heading('#5', text='Path', anchor=W)
    tree.grid(row=1, column=2, sticky = "nsew")
    if debug: print("in PopulateTree")
    treeScroll = Scrollbar(orient=VERTICAL)
    treeScroll.configure(command=tree.yview)
    tree.configure(yscrollcommand=treeScroll.set)
    treeScroll.grid(row=1, column=3, sticky=N+S)
    treeXScroll = Scrollbar(orient=HORIZONTAL)
    treeXScroll.configure(command=tree.xview)
    tree.configure(xscrollcommand=treeXScroll.set)
    treeXScroll.grid(row=2, column=2, sticky=W+N+E)
    tree.bind("<Double-Button-1>", lambda eff: populateTextWidget(eff, tree, textWidget, label, label2, textSentencesDict, scalePage)) 
    tree.bind("<Button-1>", lambda eff: handleOpenEvent(eff, tree, "")) 
    tree.bind("<Button-3>", lambda eff: handleCloseEvent2(eff, tree, ""))
##  The following binding not working, neither copy, cut nor paste. 
##    tree.bind('<Button-3>',rClicker, add='')    button.bind('<ButtonRelease-1>', lambda eff: handleOpenEvent2(eff, tree, ""))
    button.bind("<ButtonRelease-1>", lambda eff: handleOpenEvent2(eff, tree, ""))
    button2.bind('<ButtonRelease-1>', lambda eff: handleCloseEvent2(eff, tree, ""))
    return tree

def get_all_children(tree, item=""):
    grandchildren = 0
    children = tree.get_children(item)
    for child in children:
##        children += tree.get_children(tree, child)
        grandchildren += len(tree.get_children(child))
    return [len(children), grandchildren]

def PopulateTree(Tree, Parent, Dictionary, Label):
    Tree["displaycolumns"]=('#0', '#1', '#2', '#3')
    global debug, debug5
    if debug5: print("Dictionary: ",Dictionary)
    global textWidget
    global info_label0_4
    Dictionary = arrangeAliases(Dictionary, debug5)
    for key in sorted(Dictionary) :
        uid = uuid.uuid4()
        if isinstance(Dictionary[key], dict):
            Tree.insert(Parent, 'end', uid, text=key, open=True)
            PopulateTree(Tree, uid, Dictionary[key], Label)
        elif isinstance(Dictionary[key], list):
            Tree.insert(Parent, 'end', uid, text=key, open=True)
            d = dict([(i, x) for i, x in enumerate(Dictionary[key])])
            PopulateTree(Tree, uid, d, Label) 
        else:
            val = Dictionary[key]
            Tree.insert(Parent, 'end', text="", values=[val[0], val[1], val[2],val[3], val[4]])
    GAC = get_all_children(Tree, item="")
    if GAC[0] > 0:
        txt = "Found "+str(GAC[1])+" named entities in "+str(GAC[0])+" categories across combined file(s)."
    else:
        txt = "No named entities found in PDF. Check whether the PDF has been OCR'd."
    Label.config(text=txt)
    return Tree

def advancePage(event, text, label, label2, scalePage):
    global textSentencesDict, fileinFocus, pageinFocus

    if not pageinFocus == len(textSentencesDict[fileinFocus]):
        pageinFocus += 1
        text.delete(1.0, END)
        pageText= ''.join(textSentencesDict[fileinFocus][pageinFocus-1])
        text.insert(INSERT,pageText)
        label.config(text="PageViewer.\n"+str(fileinFocus))
        
        label2.config(text="Page "+str(pageinFocus)+" of "+str(len(textSentencesDict[fileinFocus])))
    else:
        label2.config(text="Reached end of file")
    scalePage.set(pageinFocus)

def ResetPage(event, scalePage, text, label, label2):
    global textSentencesDict, fileinFocus, pageinFocus
    
    pageinFocus = round(scalePage.get())
    print(pageinFocus)
    text.delete(1.0, END)
    pageText= ''.join(textSentencesDict[fileinFocus][pageinFocus-1])
    text.insert(INSERT,pageText)
    label.config(text="PageViewer.\n"+str(fileinFocus))
        
    label2.config(text="Page "+str(pageinFocus)+" of "+str(len(textSentencesDict[fileinFocus])))

##    scalePage.grid(row=4, column=4, pady=5, sticky = E+W+N)
        
def retreatPage(event, text, label, label2, scalePage):
    global textSentencesDict, fileinFocus, pageinFocus
    
    if not pageinFocus == 1:
        pageinFocus -= 1
        text.delete(1.0, END)
        pageText= ''.join(textSentencesDict[fileinFocus][pageinFocus-1])
        text.insert(INSERT,pageText)
##        
##        text.insert(INSERT,textSentencesDict[fileinFocus][pageinFocus-1])
        label.config(text="PageViewer.\n"+str(fileinFocus))
        
        label2.config(text="Page "+str(pageinFocus)+" of "+str(len(textSentencesDict[fileinFocus])))
    else:
        label2.config(text="Reached start of file")
    scalePage.set(pageinFocus)

def populateTextWidget(event, tree, textWidget, label, label2, textSentencesDict, scalePage):
##def populateTextWidget(eff=None, label=object):
    global debug, results, pageinFocus, fileinFocus, InvColorDictLabelstoColors
    item_id = event.widget.focus()
    item = event.widget.item(item_id)

    try:
        values = item['values']
##        print(values)
    except:
        label.config(text="PageViewer. No named entity instance selected. \nTry again by double-clicking in the 'File', 'Page', 'Sentence' or 'Context'\ncolumns in the Analysis window.")
        return
    parent_iid = tree.parent(item_id)
    parent_item = tree.item(parent_iid)
    key_iid = tree.parent(parent_iid)
    key = tree.item(key_iid)
    if not key["text"] in InvColorDictLabelstoColors:
        grandparent_iid = tree.parent(key_iid)
        grandparent_item = tree.item(key_iid)
        if not grandparent_item["text"] in InvColorDictLabelstoColors:
            greatgrandparent_iid = tree.parent(grandparent_iid)
            greatgrandparent_item = tree.item(grandparent_iid)
    fileinFocus = str(values[4])
    pageinFocus = values[1]
    sentenceinFocus = values[2]
    context = str(values[3])
    sentences = textSentencesDict[fileinFocus][pageinFocus-1]
    pageText= ''.join(sentences)
    block0 = ""
    string = str(parent_item["text"])
    textWidget.delete(1.0, END)

    for i in range(len(sentences)):
        if i == sentenceinFocus-1:

            part = sentences[i].partition(string)

            textWidget.insert(INSERT,part[0], 'LO')
            textWidget.insert(INSERT,part[1], 'HI')
            pos = textWidget.index(INSERT)
            textWidget.insert(INSERT,part[2], 'LO')
        else:
            textWidget.insert(INSERT,sentences[i], 'LO')
            
    try:
        color = InvColorDictLabelstoColors[key["text"]][0]
    except:
        try:
            color = InvColorDictLabelstoColors[grandparent_item["text"]][0]
        except:
            color = InvColorDictLabelstoColors[greatgrandparent_item["text"]][0]
    textWidget.tag_configure('HI', background = rgb2hex(color))
    textWidget.tag_configure('HI', font = ("Courier", 9, "bold"))
    textWidget.tag_configure('LO', font = ("Courier", 8))
    textWidget.see(pos)
    textWidget.bind('<Button-3>',rClicker, add='')
    
    label.config(text="PageViewer.\n"+str(fileinFocus))
    pages = len(textSentencesDict[fileinFocus])    
    label2.config(text="Page "+str(pageinFocus)+" of "+str(pages))

    scalePage.config(orient=HORIZONTAL, resolution=1, length = 300)
    scalePage.config(from_=1, to=pages, showvalue = 1, tickinterval=round(pages/5))
    scalePage.set(pageinFocus)
    scalePage.grid(row=4, column=4, pady=5)
    scalePage.bind('<ButtonRelease-1>', lambda eff: ResetPage(None, scalePage, textWidget, label, label2))

def rgb2hex(rgb):
    return "#%02x%02x%02x" % rgb1to255(rgb)

def rgb1to255(tpl):
    return (round(tpl[0]*255),round(tpl[1]*255), round(tpl[2]*255))

def rClicker(e):
    ''' right click context menu for all Tk Entry and Text widgets
    '''
    global fileinFocus, pageinFocus, textWidget
    try:
        def rClick_Copy(e, apnd=0):
##            chars = textWidget.get("sel.first", "sel.last")
## Need to build an insert of metadata mechanism
##            pyperclip.copy('File: '+fileinFocus+"Page: "+str(pageinFocus)+"\n|"+chars)
            e.widget.event_generate('<Control-c>')

        def rClick_Cut(e):
            e.widget.event_generate('<Control-x>')

        def rClick_Paste(e):
            e.widget.event_generate('<Control-v>')

        e.widget.focus()

        nclst=[
               (' Cut', lambda e=e: rClick_Cut(e)),
               (' Copy', lambda e=e: rClick_Copy(e)),
               (' Paste', lambda e=e: rClick_Paste(e)),
               ]

        rmenu = Menu(None, tearoff=0, takefocus=0)

        for (txt, cmd) in nclst:
            rmenu.add_command(label=txt, command=cmd)

        rmenu.tk_popup(e.x_root+40, e.y_root+10,entry="0")

    except TclError:
        print (' - rClick menu, something wrong')
        pass

    return "break"

def rClickbinder(r):

    try:
        for b in [ 'Text', 'Entry', 'Listbox', 'Label']: #
            r.bind_class(b, sequence='<Button-3>',
                         func=rClicker, add='')
    except TclError:
        print (' - rClickbinder, something wrong')
        pass

def ExtractHighlights(eff = None, listbox= object, label = object):
    global time, file_sel_list, debug
    file_sel_list = get_list(None, listbox, label)
    if debug: print("in ExtractHighlights")
    timestr = time.strftime("%Y%m%d-%H%M%S")
##    print(file_sel_list)
    successText, failureText, failure2Text = "", "", ""
##    try:
    for file in file_sel_list:
        outputFileName = file+timestr[9:]+".xlsx"
        res = pdfannot2df(file, outputFileName, debug)
        if res[0] != "": successText += ", "+res[0]
        if res[1] != "": failureText += "\n"+res[1]
        if res[2] != "": failure2Text += "\n"+res[2]
    label.config(text="Highlighted file(s) extracted with filename(s) ending as ending in '"+timestr+".xlsx'.\n"+"The following file(s) contained no readable highlights:"+ failureText+ failure2Text+" .", justify = LEFT)
    
##        label.config(text="Saved with filename"+outputFileName)
##    except:
##        label.config(text="Export unsuccessful")
    
def pdfannot2df(input_pdf, outputFileName, debug):
    """Takes an annotated pdf as an input and transforms it into a dlf
    :param input_pdf: path to the pdf.
    :return:the adf corresponding to the pdf's annotations
    """

    pdf = fitz.open(input_pdf)
    l, order = [], 0
    context = ""
    contentDict = {}
    for ixpage, page in enumerate(pdf):
        print (page)
        tmp = {'Page': ixpage + 1, 'File': input_pdf}
##        tmp = {'page': ixpage + 1, 'pdf_path': input_pdf, 'page_width': page.rect[2], 'page_height': page.rect[3]}
        words = page.getTextWords()
        annot = page.firstAnnot
        print('annot : ', annot) if debug else 0
        print('page : ', ixpage) if debug else 0
        while annot:
            print('type annot : ', annot.type[1]) if debug else 0
            mywords = []
            date1, date2, content, context = "", "", "", ""
            
            if annot.type[1] == 'Highlight':
                mywords, annot = _extract_word_from_highlight(annot, words)

            elif annot.type[1] == 'Squiggly':
                mywords, annot = _extract_word_from_highlight(annot, words)

            elif annot.type[1] == 'Underline':
                mywords, annot = _extract_word_from_highlight(annot, words)

            elif annot.type[1] == 'Square':
                try:
                    mywords = [w for w in words if fitz.Rect(w[:4]).intersects(annot.rect)]
                except:
                    mywords = ""

# helpful to build functionality here that would concatenate results from closely adjacent annotations of the same type - trouble is that 'closely adjacent' will depend on the text size

            else:
                print("Encountered an unknown annotation type.") if debug else 0

            annot_text = " ".join(w[4] for w in mywords)

            print(mywords) if debug else 0

            order += 1
            print('order : ', order) if debug else 0

# unpack dict in annot.info['content'] field

##            content = annot.info['content']
##            
##            print("content: "+content)if debug else 0
##
####            json_acceptable_content = content.replace("'", "\"")
####            contentDict = json.loads(json_acceptable_content)
##
##            if content != "":
##                print ("in content loop")if debug else 0
##                try:
##                    contentDict = eval(content)
##                    print("contentDict: ",contentDict)if debug else 0
##                    if contentDict["date"]["values"][0]["type"] == "daterange":
##                        date1 = contentDict["date"]["values"][0]["start"]
##                        date2 = contentDict["date"]["values"][0]["end"]
##                    elif contentDict["date"]["values"][0]["type"] == "date":
##                        date1 = contentDict["date"]["values"][0]["value"]
##                    context = contentDict["context"]
##                except:
##                    print ("exception")if debug else 0
##                
##  are there any other types of dates in the MS_Recognizer output? 

# get RGB triple of annot color 
            
            lst = (round(annot.colors['stroke'][0],2),round(annot.colors['stroke'][1],2),
                                                                               round(annot.colors['stroke'][2],2))
            colorKey = InvColorDicttoLabels[lst]
##            print("colorKey[:6]",colorKey[:6])
            if colorKey[:6] == "Manual":
                colorKey = "HIGHLIGHTS"
##            print(lst)
##            print(InvColorDicttoLabels)
                
# update tmp dictionary
            
            if (lst) in InvColorDicttoLabels:
                tmp.update({\
##                    'x': int(annot.rect[0]), 'y': int(annot.rect[1]),
##                        # Those might be wrong for multi line highlights as the rect only
##                        # correspond to the one of the last line
##                        'w': int(annot.rect[2] - annot.rect[0]), 'h': int(annot.rect[3] - annot.rect[1]),\
                            'Term': annot_text, \
                            'Sentence':0,\
                            'Context':"",\
                            'Note': annot.info['content'], \
                            'Type': annot.type[1], \
                            'color':lst, \
                            'Hyperlink':input_pdf,

##                            'color':annot.colors['stroke'],
##                            'ColorName':InvColorDicttoNames[lst], \
                            'colorKey':colorKey,\
                            'order': order,\
##                            'date1': date1, 'date2': date2, 'context':context\
                            })
            else:
                tmp.update({\
##                    'x': int(annot.rect[0]), 'y': int(annot.rect[1]),
##                        # Those might be wrong for multi line highlights as the rect only
##                        # correspond to the one of the last line
##                        'w': int(annot.rect[2] - annot.rect[0]), 'h': int(annot.rect[3] - annot.rect[1]),\
                            'Term': annot_text, \
                            'Sentence':0,\
                            'Context':"",\
                            'Note': annot.info['content'], \
                            'Type': annot.type[1], \
                            'color':lst, \
                            'Hyperlink':input_pdf,

##                            'color':annot.colors['stroke'],
##                            'ColorName':"Unknown", \
                            'colorKey':"Unknown",\
                            'order': order,\
##                            'date1': date1, 'date2': date2, 'context':context\
                            })
##                tmp.update({\
##                    'x': int(annot.rect[0]), 'y': int(annot.rect[1]),'w': int(annot.rect[2] - annot.rect[0]), 'h': int(annot.rect[3] - annot.rect[1]),\
##                    'type': annot.type[1], \
##                            'note': annot.info['content'], 'color':annot.colors['stroke'],'colorName':"Unknown", 'colorKey':"Unknown",'order': order, 'text': annot_text, 'date1': date1, 'date2': date2, 'context':context})
##            print (tmp)
            print(tmp) if debug else 0
            l.append(deepcopy(tmp))
            annot = annot.next

    print("l is: "+str(l))
    adf = pandas.DataFrame(l)
    print('adf : ', adf) if debug else 0
    successText, failureText, failure2Text = "", "", ""
    
    if adf.empty :
        print(f'The document {input_pdf} does not contain any annotations.')
        failureText = input_pdf

##    elif True:
##        try:
##            boo = adf[adf.get('type').isnull()].shape[0]
##            if adf[adf.get('type').isnull()].shape[0]:
##                raise Exception(f'Missing {adf[adf["type"].isnull()].shape[0]} type annotation(s) in {input_pdf}')
##                failure2Text = input_pdf
##         except:
##            0
        
    else :
        
## need to capture and insert line number (or sentence number)

        final_columns = ["Term","File","Page","Sentence","Context","Note", "Type","color","colorKey","Hyperlink"]
##        final_columns = ['page', 'type', 'text', 'note', 'color', 'colorName', 'colorKey', 'date1', 'date2']
##        final_columns = ['order', 'page', 'x', 'y', 'w', 'h', 'type', 'label', 'color', 'colorName', 'colorKey', 'page_height', 'page_width','pdf_path', 'text', 'date1', 'date2', 'context']
        adf = adf[final_columns]

##Export each group of highlights to a separate sheet
        colorKeys = adf['colorKey'].unique().tolist()
        if debug: print("165")
        writer = pandas.ExcelWriter(outputFileName, engine='xlsxwriter')
        for color in colorKeys:
            mydf = adf.loc[adf.colorKey==color]
            mydf.to_excel(writer, sheet_name=color)

        writer.save()
        successText = outputFileName
        
    return successText, failureText, failure2Text

def _extract_word_from_highlight(annot, words, debug=False):
    """Extracts words behind a highlight
    :param annot: a highlight annotation to extract words from. Warning : if it is followed by others
    highlight annotations which labels are named specifically ('same_label'+ -/- 'next_integer'),
     their words will be extracted too.
     :param words: the words of the page containing the annot.
     :param debug: debug param
    :return: the list of words extracted and the annotation which may have changed.
    """

    mywords = []

    #  If the highlight spans on multiple text boxes (possibly multiple lines
    if len(annot.vertices) > 4:

        for k in range(len(annot.vertices) // 4):
            rectangle = fitz.Rect(annot.vertices[k * 4][0], annot.vertices[k * 4][1],
                                  annot.vertices[k * 4 + 3][0], annot.vertices[k * 4 + 3][1])

            for w in words:
                r = deepcopy(rectangle)
                area_intersect = r.intersect(fitz.Rect(w[:4])).getRectArea()
                area_word = fitz.Rect(w[:4]).getRectArea()

                if area_intersect / area_word > 0.6:
                    mywords.append(w)

    else:
        mywords += [w for w in words if
                    annot.rect.intersect(fitz.Rect(w[:4])).getRectArea() / fitz.Rect(w[:4]).getRectArea() > 0.6]

        if annot.next and annot.next.info['content']:
            label_next = annot.next.info['content']
            print(annot.next.info) if debug else 0

            while annot.next and annot.next.info['content'] and label_next[(len(label_next) - 4):(len(label_next) - 1)] == '-/-' and int(label_next[len(label_next) - 1]) > 1:
                annot = annot.next
                mywords += [w for w in words if
                            annot.rect.intersect(fitz.Rect(w[:4])).getRectArea() / fitz.Rect(w[:4]).getRectArea() > 0.6]
                if annot.next and annot.next.info['content']:
                    label_next = annot.next.info['content']

    return mywords, annot


##and do the following to display each PDF page image:

#-----------------------------------------------------------------
# MuPDF code
def PreviewPDF(event, root, listbox):

    global debug, results, pageinFocus, fileinFocus, InvColorDictLabelstoColors
    
    file_sel_list = get_list(None, listbox, None)

##    def do_zoom(event):
##        factor = 1.001 ** event.delta
##        canvas.scale(ALL, event.x, event.y, factor, factor)

##    populateTextWidget(eff, tree, textWidget, label, label2, textSentencesDict, scalePage)
    top = Toplevel()

    top.title("PDF Viewer...")

    x = root.winfo_x()
    y = root.winfo_y()
    top.geometry("+%d+%d" % (x + 50, y + 0))

# zoom factor
    zoom = 0.9

    mat = fitz.Matrix(zoom, zoom)

    print("fileinFocus:", fileinFocus)
##    doc = fitz.open(fileinFocus)
#DEBUGGING HERE
##    print("file_sel_list: ", file_sel_list)

##    Code from PDFexport
##    try:
##        for file in file_sel_list:
##    ######             print ("254")
##            d2 = results[tuple(file_sel_list)][2]
##    ##            pprint(d2)
##            p0 = os.path.split(file)[0]
##            p1 = os.path.split(file)[1]
##    ##            doc = markup(p1, d2[file], InvColorDictLabelstoColors, debug7)
##            doc = markup(file, d2[file], InvColorDictLabelstoColors, debug6)
##    ##            saveDocToPDF(doc, p0, p1, outfldr, timestr, debug)
##            pathnew = p0+pathsep+outfldr+pathsep+p1[:-4]+" "+timestr+" HI.pdf"
##            doc.save(pathnew)
##            doc.close()
##        label.config(text="Saved with filename(s) ending in '"+timestr+"_HI.pdf'")
##    except:
##        label.config(text="Export failed, please try other file(s)"+str(file_sel_list))

    d2 = results[tuple(file_sel_list)][2]
    
##    print(d2)
    ##            doc = markup(p1, d2[file], InvColorDictLabelstoColors, debug7)
    

    p1 = os.path.split(fileinFocus)[1]
    doc = markup(p1, d2[fileinFocus], InvColorDictLabelstoColors, debug6)
    
    pix = doc[pageinFocus-1].getPixmap(matrix = mat)
##    pix = doc.getPagePixmap(0)
    mode = "RGBA" if pix.alpha else "RGB"
    img = PIL.Image.frombytes(mode, [pix.width, pix.height], pix.samples)
    canvas = Canvas(top, width=pix.width, height=pix.height)
    canvas.grid(row=0, column=0, sticky = "nsew")
    canvas.image = PIL.ImageTk.PhotoImage(img)
    canvasPdfs = canvas.create_image(0, 0, anchor=NW, image=canvas.image)

##    canvas.bind("<MouseWheel>", do_zoom)
##    canvas.bind('<ButtonPress-1>', lambda event: canvas.scan_mark(event.x, event.y))
##    canvas.bind("<B1-Motion>", lambda event: canvas.scan_dragto(event.x, event.y, gain=1))

##    currentPage = pageinFocus - 1#-----------------------------------------------------------------

    def generatePic(zoom, doc):

        mat = fitz.Matrix(zoom, zoom)
        pix = doc[pageinFocus-1].getPixmap(matrix = mat)
    ##    pix = doc.getPagePixmap(0)
        mode = "RGBA" if pix.alpha else "RGB"
        img = PIL.Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        canvas = Canvas(top, width=pix.width, height=pix.height)
        canvas.grid(row=0, column=0, sticky = "nsew")
        canvas.image = PIL.ImageTk.PhotoImage(img)
        canvasPdfs = canvas.create_image(0, 0, anchor=NW, image=canvas.image)
##        canvas.bind("<MouseWheel>", do_zoom)
##        canvas.bind('<ButtonPress-1>', lambda event: canvas.scan_mark(event.x, event.y))
##        canvas.bind("<B1-Motion>", lambda event: canvas.scan_dragto(event.x, event.y, gain=1))

    generatePic(zoom, doc)

    def nxtBtn_Click(event, canvas, canvasPdfs, zoom, doc):
        
        global pageinFocus
        
        canvas.delete(canvasPdfs)

        pageinFocus += 1

        generatePic(zoom, doc)

##        advancePage(None, textWidget, \
##            info_label0_4, info_label_2_4, scalePage)

    def prvBtn_Click(event, canvas, canvasPdfs, zoom, doc):
        
        global pageinFocus
        
        canvas.delete(canvasPdfs)

        pageinFocus -= 1

        generatePic(zoom, doc)

    def inBtn_Click(event, canvas, canvasPdfs, zoom, doc):
        
        zoom *= math.sqrt(2)
        
        canvas.delete(canvasPdfs)

        generatePic(zoom, doc)

        
    def outBtn_Click(event, canvas, canvasPdfs, zoom, doc):
        
        zoom *= math.sqrt(0.5)
        
        canvas.delete(canvasPdfs)

        generatePic(zoom, doc)

    PrevBtn = Button(top, text="Prev")
    PrevBtn.grid(column=1, row=0, sticky = N)
    PrevBtn.bind("<Button-1>", lambda eff: prvBtn_Click(eff, canvas, canvasPdfs, zoom, doc))

    nxtBtn = Button(top, text="Next")
    nxtBtn.grid(column=2, row=0, sticky = N)
    nxtBtn.bind("<Button-1>", lambda eff: nxtBtn_Click(eff, canvas, canvasPdfs, zoom, doc))

    inBtn = Button(top, text="Zoom +")
    inBtn.grid(column=1, row=0, sticky = S)
    inBtn.bind("<Button-1>", lambda eff: inBtn_Click(eff, canvas, canvasPdfs, zoom, doc))

    outBtn = Button(top, text="Zoom -")
    outBtn.grid(column=2, row=0, sticky = S)
    outBtn.bind("<Button-1>", lambda eff: outBtn_Click(eff, canvas, canvasPdfs, zoom, doc))

    top.mainloop()

def get_text_percentage(file_name: str) -> float:
    """
    Calculate the percentage of document that is covered by (searchable) text.

    If the returned percentage of text is very low, the document is
    most likely a non-searchable PDF
    """
    total_page_area = 0.0
    total_text_area = 0.0

    doc = fitz.open(file_name)

    for page_num, page in enumerate(doc):
        total_page_area = total_page_area + abs(page.rect)
        text_area = 0.0
        for b in page.getTextBlocks():
            r = fitz.Rect(b[:4])  # rectangle where block text appears
            text_area = text_area + abs(r)
        total_text_area = total_text_area + text_area
    doc.close()
    return total_text_area / total_page_area

def ocr_pdf_if_not_searchable(eff = None, listbox = object, label = object):
##    print("flag1")
    global file_sel_list
    filesProcessed = ""
    x = 0
    file_sel_list = get_list(None, listbox, label)
    for file in file_sel_list:
        head, tail = os.path.split(file)
        text_perc = get_text_percentage(file)
        if text_perc < 0.01:
            x +=1
            result = ocrmypdf.ocr(file, file[:-4]+"_OCR.pdf", redo_ocr = True)
            filesProcessed += tail
    label.config(text=("Total of "+str(x)+" files saved with '_OCR' suffix."))

           
class pdfAnalyst(object):

    global file_list, InvColorDictLabelstoColors

    def __init__(self, parent):

        root.columnconfigure(0, weight=1, minsize=200)
        root.columnconfigure(1, weight=0)
        root.columnconfigure(2, weight=2, minsize=600)
        root.columnconfigure(3, weight=0)
        root.columnconfigure(4, weight=3, minsize=600)
        root.columnconfigure(5, weight=0)
        
        root.rowconfigure(0, weight=0)
        root.rowconfigure(2, weight=0)
        root.rowconfigure(1, weight=1)
        root.rowconfigure(2, weight=0)
        root.rowconfigure(3, weight=0)
        root.rowconfigure(4, weight=0)
        root.rowconfigure(5, weight=0)
        root.rowconfigure(6, weight=0)
        root.rowconfigure(7, weight=0)

        info_label0_0 = Label(fg='red')
        info_label0_0.grid(row=0, column=0, pady=5, sticky=E+W)
        info_label0_0.config(text="Sandbox\nDouble-click a file to analyse it or\nchoose an option below.")

        info_label0_2 = Label(fg='red')
        info_label0_2.grid(row=0, column=2, pady=5, sticky=E+W)
        info_label0_2.config(text="Analysis results\nDouble-click a term to display it in the Pageviewer or choose an export option below.")
        
        info_label0_4 = Label(fg='red')
        info_label0_4.grid(row=0, column=4, pady=5, sticky=E+W)
        info_label0_4.config(text="Pageviewer")

        # selectmode=EXTENDED allows ctrl/shift mouse slections
        listbox1 = Listbox(root, width=40, height=30, selectmode=EXTENDED)
        listbox1.grid(row=1, column=0, sticky = "nsew")
        listbox1.bind('<Double-1>', lambda eff: analyse_file(None, listbox1, textWidget, info_label_2_0, info_label0_4, info_label_2_2, info_label_2_4, tree, root, scalePage, button_4_2_1, button_4_2_1b))

        yscroll1 = Scrollbar(command=listbox1.yview, orient=VERTICAL)
        yscroll1.grid(row=1, column=1, sticky=N+S+W)
        listbox1.configure(yscrollcommand=yscroll1.set)

        xscroll1 = Scrollbar(command=listbox1.xview, orient=HORIZONTAL)
        xscroll1.grid(row=2, column=0, sticky=E+N+W)
        listbox1.configure(xscrollcommand=xscroll1.set)

        textWidget = Text(root, width=40, height=30, wrap = WORD)
        textWidget.grid(row=1, column=4, sticky = "nsew")

        button_4_2_1 = Button(root, text='Expand all', command=None)
        button_4_2_1.grid(row=5, column=2, pady=5, sticky = W)

        button_4_2_1b = Button(root, text='Collapse all', command=None)
        button_4_2_1b.grid(row=4, column=2, pady=5, sticky = W)

        yscroll3 = Scrollbar(command=textWidget.yview, orient=VERTICAL)
        yscroll3.grid(row=1, column=5, sticky=N+S+W)
        textWidget.configure(yscrollcommand=yscroll3.set)
        
        xscroll3 = Scrollbar(command=textWidget.xview, orient=HORIZONTAL)
        xscroll3.grid(row=2, column=4, sticky=E+N+W)
        textWidget.configure(xscrollcommand=xscroll3.set)
        
        info_label_2_0 = Label(fg='red')
        info_label_2_0.grid(row=3, column=0, pady=5, sticky=W)
        
        button_3_0_0 = Button(root, text='Sort file list')
        button_3_0_0.grid(row=4, column=0, pady=5, sticky = W)
        button_3_0_0.bind('<ButtonRelease-1>', lambda eff: sort_list(eff, listbox1))

        button_3_0_2 = Button(root, text='Analyse multiple file(s)')
        button_3_0_2.grid(row=4, column=0, pady=5, sticky = E)
        button_3_0_2.bind('<ButtonRelease-1>', lambda eff: analyse_file(None, listbox1, textWidget, info_label_2_0, info_label0_4, info_label_2_2, info_label_2_4, tree, root, scalePage, button_4_2_1, button_4_2_1b))
        
        scalePage = Scale(root, orient=HORIZONTAL, from_=0, to=1, length = 300, showvalue = True)
        scalePage.grid(row=4, column=4, pady=5)
        scalePage.set(pageinFocus)
        scalePage.bind('<ButtonRelease-1>', lambda eff: ResetPage(None, scalePage, textWidget, info_label0_4))
        
        button_3_4_0 = Button(root, text='<< Previous')
        button_3_4_0.grid(row=4, column=4, pady=5, sticky = W)
        button_3_4_0.bind('<ButtonRelease-1>', lambda eff: retreatPage(None, textWidget, info_label0_4, info_label_2_4, scalePage))

        button_3_0_2 = Button(root, text='Next >>')
        button_3_0_2.grid(row=4, column=4, pady=5, sticky = E)
        button_3_0_2.bind('<ButtonRelease-1>', lambda eff: advancePage(None, textWidget, \
                                                                       info_label0_4, info_label_2_4, scalePage))

        info_label_2_2 = Label(fg='red')
        info_label_2_2.grid(row=3, column=2, pady=5, ipady=5, sticky=W)
        info_label_2_2.config(text="Placeholder for info_label_2_2", wraplength=600)

        info_label_2_4 = Label(fg='red')
        info_label_2_4.grid(row=3, column=4, pady=5, ipady=5, sticky=E+W)
        info_label_2_4.config(text="Page navigation")

        button_5_0_0 = Button(root, text='Import')
        button_5_0_0.grid(row = 5, column=0, pady=5, stick = W)
        button_5_0_0.bind('<ButtonRelease-1>', lambda eff: import_files(eff, listbox1, info_label_2_0))

        button_5_0_1 = Button(root, text='Remove')
        button_5_0_1.grid(row = 5, column=0, pady=5)
        button_5_0_1.bind('<ButtonRelease-1>', lambda eff: remove_files(eff, listbox1, info_label_2_0))

        button_5_0_2 = Button(root, text='Remove All')
        button_5_0_2.grid(row = 5, column=0, pady=5, sticky=E)
        button_5_0_2.bind('<ButtonRelease-1>', lambda eff: remove_all_files(eff, listbox1, info_label_2_0))

        button_4_2_0 = Button(root, text='Export summary analysis', command=None)
        button_4_2_0.grid(row=4, column=2, pady=5, sticky = E)
        button_4_2_0.bind('<ButtonRelease-1>', lambda eff: ExportDicttoExcelUVO(eff, info_label_2_2))

        tree = CreateTree(root, textWidget, info_label0_4, info_label_2_2, textSentencesDict, scalePage, button_4_2_1, button_4_2_1b)
        tree = PopulateTree(tree, '', result, info_label_2_2)

        button_5_2_0 = Button(root, text='Export detailed analysis', command=None)
        button_5_2_0.grid(row = 5, column=2, pady=5, sticky = E)
        button_5_2_0.bind('<ButtonRelease-1>', lambda eff: ExportDetailstoExcel(eff, info_label_2_2))

        button_4_2_1c = Button(root, text='Delete existing highlights', command=None)
        button_4_2_1c.grid(row=4, column=2, pady=5)
        button_4_2_1c.bind('<ButtonRelease-1>', lambda eff: DeleteHighlights(eff, listbox1, info_label_2_2))
       
        button_4_2_2 = Button(root, text='Extract existing highlights', command=None)
        button_4_2_2.grid(row=5, column=2, pady=5)
        button_4_2_2.bind('<ButtonRelease-1>', lambda eff: ExtractHighlights(eff, listbox1, info_label_2_2))

        button_5_4_2 = Button(root, text='Preview PDF', command=None)
        button_5_4_2.grid(row=5, column=4, pady=5, sticky = E)
        button_5_4_2.bind('<ButtonRelease-1>', lambda eff: PreviewPDF(eff, root, listbox1))

        button_5_2_1 = Button(root, text='Export Highlighted PDF(s)', command=None)
        button_5_2_1.grid(row = 5, column=4, pady=5, sticky = W)
        button_5_2_1.bind('<ButtonRelease-1>', lambda eff: ExporttoPDF(eff, info_label_2_2))

        button_6_4_1 = Button(root, text='OCR non-searchable PDF(s)', command=None)
        button_6_4_1.grid(row = 6, column=4, pady=5, sticky = W)
        button_6_4_1.bind('<ButtonRelease-1>', lambda eff: ocr_pdf_if_not_searchable(eff, listbox1, info_label_2_2))

        button_5_2_2 = Button(root, text='Customize highlights', command = None)
        button_5_2_2.grid(row = 5, column=4, pady=5)
        button_5_2_2['command'] = lambda: CustomizeHL(False)

        button_6_2_1 = Button(root, text='Extract images', command=None)
        button_6_2_1.grid(row=6, column=2, pady=5, sticky = W)
        button_6_2_1.bind('<ButtonRelease-1>', lambda eff: ExtractImages(eff, listbox1, info_label_2_2))

        button_6_2_2 = Button(root, text='Export Chronology to Word', command=None)
        button_6_2_2.grid(row=6, column=2, pady=5)
        button_6_2_2.bind('<ButtonRelease-1>', lambda eff: Export2Word(eff, info_label_2_2))

        button_6_2_3 = Button(root, text='Save PDF pages to images', command=None)
        button_6_2_3.grid(row=6, column=2, pady=5, sticky = E)
        button_6_2_3.bind('<ButtonRelease-1>', lambda eff: PDF2Images(eff, listbox1, info_label_2_2))
       

        info_label2 = Label(fg='red')
        info_label2.grid(row=2, column=2, pady=5, sticky=W)

        file_list = import_files(None, listbox1, info_label_2_0)
     
        root.update_idletasks()
        root.minsize(800, 800)
        ent = Entry(root, width=50)
        ent.bind('<Button-3>',rClicker, add='')
        root.mainloop()

if __name__ == "__main__":
    root = Tk()
    root.state('zoomed')
    root.title("pdfAnalyst")
    root.iconbitmap("ico.bmp")
    myapp = pdfAnalyst(root)
    root.mainloop()
